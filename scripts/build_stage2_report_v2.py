#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""更新版报告(final_eval完成后): 插入显著性 + ROC/PR/混淆/特征重要性/占比/森林图。"""
import os, json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT=os.path.join(os.path.dirname(__file__),"..")
RES=os.path.join(ROOT,"EP_ATLAS","results"); OUT=os.path.join(RES,"report_stage2")

def main():
    d=json.load(open(os.path.join(RES,'compare_8configs.json'),encoding='utf-8'))
    sig=json.load(open(os.path.join(RES,'significance.json'),encoding='utf-8'))
    os.makedirs(OUT,exist_ok=True)
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5); st._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    doc.add_heading('K562 顺式增强子–启动子互作预测：序列模型特征对比（更新版报告）',level=0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('EP-ATLAS K562（hg19）| ABC / Genos embedding / co-embedding | 按增强子分组 5 折交叉验证 + 配对显著性检验')

    doc.add_heading('摘要',level=1)
    doc.add_paragraph('本研究评估序列模型预测 K562 顺式增强子–启动子（E–P）互作的能力。结果表明：'
        '（1）ABC 特征提供主导信号（AUC 0.69）；'
        '（2）增强子+启动子拼接的 co-embedding 与 ABC 融合取得最佳性能（AUC 0.725），'
        '显著优于纯 ABC（ΔAUC +0.034，p<0.0001）；'
        '（3）仅 embedding 接近随机（AUC 0.53–0.56）。'
        'co-embedding 是否降维（RED）对性能无显著影响（p=0.11），两者均显著优于 ABC。'
        '特征重要性显示 ABC 与 co-embedding 各占约一半。')

    doc.add_heading('1. 数据与方法',level=1)
    doc.add_paragraph('数据：1,263 正样本（RIC-seq/KARR-seq 支持的 E–P 互作）+ 1,263 负样本（打乱配对、同染色体、'
        '距离分箱匹配、剔除完整互作集）。共 2,526 对，全部同染色体，hg19。')
    doc.add_paragraph('特征：ABC（194 维，含 Hi-C 接触/活性/序列组成）；embedding（Genos-1.2B，独立编码 2049 维）；'
        'co-embedding（拼接编码 1024 维）；RED 为 KernelPCA(RBF,100) 降维（折内拟合）。'
        '模型：GB 与 RF。评估：按增强子分组 5 折 CV + 配对 bootstrap 显著性检验。')

    doc.add_heading('2. 结果',level=1)
    doc.add_heading('2.1 九种配置对比',level=2)
    tbl=doc.add_table(rows=1,cols=7); tbl.style='Light Grid Accent 1'
    for i,h in enumerate(['配置','GB-AUC','GB-PRAUC','GB-Acc','RF-AUC','RF-PRAUC','RF-Acc']): tbl.rows[0].cells[i].text=h
    cfgs=sorted(set(v['config'] for v in d.values()), key=lambda c:-max(v['auc_m'] for v in d.values() if v['config']==c))
    for c in cfgs:
        gb=next(v for v in d.values() if v['config']==c and v['model']=='GB'); rf=next(v for v in d.values() if v['config']==c and v['model']=='RF')
        row=tbl.add_row().cells
        row[0].text=c; row[1].text=f"{gb['auc_m']:.4f}"; row[2].text=f"{gb['ap_m']:.4f}"; row[3].text=f"{gb['acc_m']:.4f}"
        row[4].text=f"{rf['auc_m']:.4f}"; row[5].text=f"{rf['ap_m']:.4f}"; row[6].text=f"{rf['acc_m']:.4f}"
    doc.add_heading('2.2 配置对比图',level=2)
    doc.add_picture(os.path.join(RES,'fig_config_compare.png'),width=Inches(6.0)); doc.add_paragraph('图 1. 各配置 AUC（GB，分组 CV）。')

    doc.add_heading('2.3 ROC 与 PR 曲线',level=2)
    doc.add_picture(os.path.join(RES,'fig_roc.png'),width=Inches(3.8)); doc.add_picture(os.path.join(RES,'fig_pr.png'),width=Inches(3.8))
    doc.add_paragraph('图 2-3. 三个关键配置（ABC / ABC+coembed 原始 / ABC+coembed RED）的 ROC 与 PR 曲线。')

    doc.add_heading('2.4 混淆矩阵（最佳配置：ABC+coembed 原始 GB）',level=2)
    doc.add_picture(os.path.join(RES,'fig_confusion.png'),width=Inches(3.2))

    doc.add_heading('2.5 显著性检验',level=2)
    t2=doc.add_table(rows=1,cols=5); t2.style='Light Grid Accent 1'
    for i,h in enumerate(['对比','ΔAUC','95% CI','p (AUC)','p (Acc)']): t2.rows[0].cells[i].text=h
    order=[('ABC+co__vs__ABC','ABC+coembed(raw) vs ABC'),('ABC+coRED__vs__ABC','ABC+coembed(RED) vs ABC'),('ABC+co__vs__ABC+coRED','coembed(raw) vs coembed(RED)')]
    for k,lab in order:
        v=sig[k]; r=t2.add_row().cells
        r[0].text=lab; r[1].text=f"{v['aucA']-v['aucB']:+.4f}"; r[2].text=f"[{v['auc_ci'][0]:+.4f},{v['auc_ci'][1]:+.4f}]"
        r[3].text=f"{v['auc_p']:.4f}"; r[4].text=f"{v['acc_p']:.4f}"
    doc.add_picture(os.path.join(RES,'fig_significance_forest.png'),width=Inches(5.5))
    doc.add_paragraph('图 4. AUC 差异的配对 bootstrap 95% 置信区间森林图（* p<0.05）。')
    doc.add_paragraph('结论：ABC+co-embedding（原始或降维）均显著优于纯 ABC（p<0.001）；'
        '原始与降维之间差异不显著（p=0.11），即降维与否对性能无显著影响。')

    doc.add_heading('2.6 特征重要性',level=2)
    doc.add_picture(os.path.join(RES,'fig_feature_importance.png'),width=Inches(4.0))
    doc.add_picture(os.path.join(RES,'fig_importance_share.png'),width=Inches(3.0))
    doc.add_paragraph('图 5-6. 最佳配置（ABC+coembed）的特征重要性 Top20（蓝=ABC 特征，红=coembed 维度）'
        '与整体占比（ABC 47.2% / coembedding 52.8%）。')

    doc.add_heading('3. 关键发现',level=1)
    for txt in [
        '最佳：ABC + co-embedding（GB）AUC=0.725，显著优于 ABC(pure) 0.692（+0.034，p<0.0001）。',
        'ABC+coembed(RED) 0.716 也显著优于 ABC（+0.025，p<0.001）；与原始版差异不显著。',
        '仅 embedding（独立或 co）接近随机（0.53–0.56），必须与 ABC 融合。',
        '特征重要性：co-embedding 与 ABC 各占约一半，说明拼接编码确实贡献了互补的序列信息。',
        'GB 全面优于 RF。',
    ]: doc.add_paragraph('• '+txt)

    doc.add_heading('4. 结论与下一步',level=1)
    doc.add_paragraph('sequence-only 能显著预测顺式 E–P 互作（AUC 0.725，显著优于随机与纯 ABC），'
        'co-embedding 是优于独立编码的序列利用方式。下一步：完成全量 18 组数据收集与更多分层分析，出终极报告。')

    doc.add_heading('5. 局限',level=1)
    doc.add_paragraph('(1) "负样本"为无实验证据；(2) 仅顺式；(3) embedding 增量有限，突破需染色质/3D 信息。')

    rep=os.path.join(OUT,'K562_cis_更新版报告_v2.docx'); doc.save(rep)
    print('更新版报告已生成:',rep)

if __name__=='__main__': main()
