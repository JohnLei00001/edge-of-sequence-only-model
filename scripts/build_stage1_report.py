#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""阶段1完整报告(现阶段, 用已有数据/图)。final_eval跑完后会更新此文档。"""
import os, json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT=os.path.join(os.path.dirname(__file__),"..")
RES=os.path.join(ROOT,"EP_ATLAS","results")
OUT=os.path.join(RES,"report_stage1")

def main():
    d=json.load(open(os.path.join(RES,'compare_8configs.json'),encoding='utf-8'))
    os.makedirs(OUT,exist_ok=True)
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5); st._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    doc.add_heading('K562 顺式增强子–启动子互作预测：序列模型特征对比（阶段一报告）',level=0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('数据：EP-ATLAS K562（hg19） | 特征：ABC / Genos embedding / co-embedding | 评估：按增强子分组 5 折交叉验证')

    doc.add_heading('摘要',level=1)
    doc.add_paragraph('本报告评估序列模型（sequence-only）能否预测 K562 顺式增强子–启动子（E–P）互作。'
        '比较了 9 种特征配置 × GB/RF 两种分类器在分组交叉验证下的表现。结果显示：'
        '（1）ABC 特征（基因组距离、Hi-C 接触、增强子活性、序列组成）提供主导信号（AUC 0.69）；'
        '（2）将增强子与启动子拼接为一条序列进行 co-embedding，再与 ABC 融合，取得最佳性能（AUC 0.725，'
        '较纯 ABC 提升约 0.03）；'
        '（3）仅使用 embedding（独立或 co-embedding）几乎无信号（AUC 0.53–0.56，接近随机）。'
        '显著性检验与 ROC/PR 曲线等分析仍在补充中，完成后将更新本报告。')

    doc.add_heading('1. 研究背景与目标',level=1)
    doc.add_paragraph('增强子是重要的顺式调控元件，通过染色质环在三维空间靠近靶基因启动子以激活转录。'
        '识别"哪个增强子调控哪个启动子"是理解基因调控、解读疾病变异的关键。本工作以 K562 细胞的顺式'
        '（同染色体）互作为对象，回答一个核心科学问题：**仅凭序列信息，能否以及能在多大程度上预测 E–P 互作？**'
        '具体比较了 ABC 模型特征与深度学习序列 embedding（Genos）单独及融合时的预测能力。')

    doc.add_heading('2. 数据',level=1)
    doc.add_paragraph('正样本（n=1,263）为 K562 中经 RIC-seq/KARR-seq 支持的 E–P 互作（EP-ATLAS）。'
        '负样本（n=1,263）由顺式增强子/启动子池打乱配对生成，约束同染色体并按距离分箱匹配，'
        '剔除完整互作集（91,310 条）中出现的任何配对，确保无实验证据。正负构成平衡集（2,526 对），全部同染色体。'
        '参考基因组 hg19。')

    doc.add_heading('3. 方法',level=1)
    doc.add_heading('3.1 特征',level=2)
    doc.add_paragraph('ABC 特征（194 维）：增强子/启动子各 93 维序列组成（核苷酸频率、GC、CpG、二/三核苷酸、熵、基序、长度等），'
        '加 7 维联合特征（Hi-C 接触 C(d)=d^(-0.87)e^(-d/3Mb)、多种相似度）、log10 距离与 ABC 评分（活性×接触，按基因归一）。')
    doc.add_paragraph('Embedding：增强子与启动子分别用 Genos（Genos-1.2B，均值池化，1024 维）编码，拼接余弦相似度（2049 维）。'
        'co-embedding：将增强子+分隔符+启动子拼接为一条序列，Genos 编码一次（1024 维）。'
        'RED 表示经 KernelPCA(RBF, 100 维) 降维（在折内拟合，无泄漏）。')

    doc.add_heading('3.2 模型与评估',level=2)
    doc.add_paragraph('分类器：Gradient Boosting (GB) 与 Random Forest (RF)。'
        '评估：按增强子身份分组 5 折交叉验证（同一增强子不跨折，避免泄漏）；'
        '所有降维与标准化仅在训练折拟合。报告 AUC、PR-AUC、准确率。')

    doc.add_heading('4. 结果',level=1)
    doc.add_heading('4.1 九种配置对比',level=2)
    tbl=doc.add_table(rows=1,cols=7); tbl.style='Light Grid Accent 1'
    hdr=['配置','GB-AUC','GB-PRAUC','GB-Acc','RF-AUC','RF-PRAUC','RF-Acc']
    for i,h in enumerate(hdr): tbl.rows[0].cells[i].text=h
    cfgs=sorted(set(v['config'] for v in d.values()), key=lambda c:-max(v['auc_m'] for v in d.values() if v['config']==c))
    for c in cfgs:
        gb=next(v for v in d.values() if v['config']==c and v['model']=='GB'); rf=next(v for v in d.values() if v['config']==c and v['model']=='RF')
        row=tbl.add_row().cells
        row[0].text=c; row[1].text=f"{gb['auc_m']:.4f}"; row[2].text=f"{gb['ap_m']:.4f}"; row[3].text=f"{gb['acc_m']:.4f}"
        row[4].text=f"{rf['auc_m']:.4f}"; row[5].text=f"{rf['ap_m']:.4f}"; row[6].text=f"{rf['acc_m']:.4f}"

    doc.add_heading('4.2 配置对比图',level=2)
    doc.add_picture(os.path.join(RES,'fig_config_compare_now.png'),width=Inches(6.2))
    doc.add_paragraph('图 1. 九种特征配置在 GB 与 RF 下的 AUC（分组 CV）。')
    doc.add_picture(os.path.join(RES,'fig_config_prauc.png'),width=Inches(6.2))
    doc.add_paragraph('图 2. PR-AUC 对比。')
    doc.add_picture(os.path.join(RES,'fig_config_acc.png'),width=Inches(6.2))
    doc.add_paragraph('图 3. 准确率对比。')

    doc.add_heading('4.3 关键发现',level=2)
    for txt in [
        '最佳配置：ABC + co-embedding（原始，GB），AUC=0.725，较纯 ABC（0.692）提升约 +0.033。',
        'ABC 特征是主导信号（AUC 0.69）；距离/接触/活性贡献最大。',
        '仅 embedding（独立或 co-embedding）接近随机（AUC 0.53–0.56），必须与 ABC 融合才有用。',
        '有趣的是，原始 co-embedding（0.725）略高于降维版（RED 0.712）；两者差异是否显著待显著性检验确认。',
        '有监督降维（PLS/LDA）未能超过无监督（KPCA），且成分越多越差（过拟合），说明 embedding 的可泛化信号有限。',
    ]: doc.add_paragraph('• '+txt)

    doc.add_heading('5. 显著性检验（补充中）',level=1)
    doc.add_paragraph('AUC 与 Accuracy 的配对显著性检验（bootstrap 置信区间 + 单侧 p 值）以及 ROC/PR 曲线、'
        '混淆矩阵、特征重要性等分析正在计算，完成后将插入本节并更新本报告。')

    doc.add_heading('6. 结论与下一步',level=1)
    doc.add_paragraph('现阶段结论：sequence-only 模型确实携带 E–P 互作信息（ABC 0.69，co-embedding 融合后 0.725），'
        '但主要来自距离/接触/活性，深度 embedding 的增量较小且需与 ABC 融合。'
        '下一步：完成显著性检验与全量 18 组数据收集，出终极报告；并评估 trans 臂。')

    doc.add_heading('7. 局限',level=1)
    doc.add_paragraph('(1) "负样本"指无实验证据，非证实不互作；(2) 本阶段仅顺式；(3) embedding 增量有限，'
        '更大幅提升需染色质/3D 信息。')

    rep=os.path.join(OUT,'K562_cis_阶段一报告.docx'); doc.save(rep)
    print('阶段一报告已生成:',rep)

if __name__=='__main__': main()
