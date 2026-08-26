#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""完整报告(方法学详尽版): 数据构建/embedding做法/拼接方式/设计理由/要点 + 结果。"""
import os, json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT=os.path.join(os.path.dirname(__file__),"..")
RES=os.path.join(ROOT,"EP_ATLAS","results"); OUT=os.path.join(RES,"report_complete")

def main():
    d=json.load(open(os.path.join(RES,'compare_8configs.json'),encoding='utf-8'))
    sig=json.load(open(os.path.join(RES,'significance.json'),encoding='utf-8'))
    os.makedirs(OUT,exist_ok=True)
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5); st._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    doc.add_heading('K562 顺式增强子–启动子互作预测：sequence-only 模型的构建与评估（完整报告）',level=0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('EP-ATLAS K562（hg19）| ABC 特征 + Genos embedding / co-embedding | 按增强子分组 5 折 CV + 配对显著性检验')

    doc.add_heading('摘要',level=1)
    doc.add_paragraph('本研究构建 sequence-only 模型预测 K562 顺式增强子–启动子（E–P）互作，系统比较 ABC 模型特征与'
        'Genos 深度序列 embedding（独立编码与拼接 co-embedding）单独及融合的表现。结果显示：'
        '（1）ABC 特征（基因组距离、Hi-C 接触、增强子活性、序列组成）提供主导信号（AUC 0.69）；'
        '（2）将增强子与启动子拼接为一条序列的 co-embedding 与 ABC 融合取得最佳性能（AUC 0.725），'
        '显著优于纯 ABC（+0.034，p<0.0001）；'
        '（3）仅用 embedding（独立或 co-embedding）接近随机（AUC 0.53–0.56），必须与 ABC 融合。'
        '本研究详细说明了数据构建、embedding 与拼接的具体做法、设计理由与关键要点。')

    doc.add_heading('1. 研究背景与目标',level=1)
    doc.add_paragraph('增强子通过染色质环在三维空间靠近靶基因启动子以激活转录。预测增强子–启动子（E–P）互作对'
        '基因调控注释与疾病变异解读意义重大。本工作回答的核心问题是：**仅凭序列信息，能否以及能在多大程度上预测顺式 '
        'E–P 互作？** 为此需要回答三个子问题：(a) 序列中是否携带互作信息；(b) 如何提取；(c) 不同提取方式的优劣。')

    doc.add_heading('2. 数据构建（含理由）',level=1)
    doc.add_paragraph('正样本（n=1,263）：K562 中经 RIC-seq 或 KARR-seq 实验支持的 E–P 互作（EP-ATLAS 资源），'
        '参考基因组 hg19。')
    doc.add_paragraph('负样本（n=1,263）：为避免"正样本多数为跨染色体"带来的先验偏置，我们将正样本按顺式/反式'
        '配平（各 1,263）。负样本由顺式增强子/启动子池打乱配对生成，约束为**同染色体**，并**按基因组距离分箱分层抽样**'
        '以匹配正样本的距离分布（避免模型学"距离近=互作"这种平凡规律），再**剔除完整互作集（91,310 条）中出现的任何配对**，'
        '确保负样本无实验证据。最终构成 2×2 平衡设计（顺/反 × 正/负），本报告分析顺式臂（2,526 对，全部同染色体）。')
    doc.add_paragraph('要点：负样本的"同染色体 + 距离匹配 + 剔除全集"三重约束，是本工作严谨性的基础——它迫使模型'
        '必须从序列本身而非距离或染色体配对先验去学习互作信号。')

    doc.add_heading('3. 建模方法',level=1)
    doc.add_heading('3.1 ABC 特征（194 维）',level=2)
    doc.add_paragraph('ABC（Activity-By-Contact）框架由增强子活性 × 接触强度构成。'
        '本实现包含：(1) 序列组成特征——增强子与启动子各 93 维（核苷酸频率、GC 含量、CpG/GpC 密度、香农熵、'
        '16 种二核苷酸与 64 种三核苷酸频率、TATA/CAAT 基序、同聚物、对数长度、AT/GC 偏斜）；'
        '(2) 6 维联合特征——Hi-C 接触（由距离衰减模型 C(d)=d^(-0.87)·exp(-d/3Mb) 导出）、序列余弦相似度、'
        '三核苷酸余弦相似度、GC/CpG/欧氏相似度；(3) log10 距离；(4) ABC 评分（活性 × 接触，按基因归一化）。'
        '活性分由增强子特征 PCA 后经逻辑回归校准得到。总计 93+93+6+1+1=194 维。')
    doc.add_paragraph('要点：Hi-C 接触项显式利用了同染色体增强子-启动子的基因组距离——这是顺式互作的主导信号；'
        '活性项捕捉"增强子是否活跃"。ABC 已覆盖距离、接触、活性、组成四类信息。')

    doc.add_heading('3.2 Genos embedding（具体做法）',level=2)
    doc.add_paragraph('Embedding 由 Genos（Genos-1.2B，flash 变体，**均值池化**）经 API '
        '（https://www.dcs.cloud/api/aigress/openai/dna_embedding）生成，每个序列输出 **1024 维**向量。'
        'API 一次只接受一条序列（不支持批量），这是工程约束。')
    doc.add_paragraph('独立 embedding：增强子序列、启动子序列分别编码，各得 1024 维，再拼接为 '
        '[enh_emb(1024) | prom_emb(1024) | 余弦相似度(1)] = **2049 维**。')
    doc.add_paragraph('co-embedding：将"增强子序列 + 100 个 N 分隔 + 启动子序列"**拼接为一条序列**，由 Genos '
        '编码一次，得 1024 维。')
    doc.add_paragraph('**为什么做 co-embedding**：独立编码后再拼接，两个向量互不相关，模型无法显式捕捉"增强子-启动子'
        '的交互关系"。co-embedding 让模型在**同一个上下文窗口**内同时看到两个元件，从而可能通过注意力计算二者的关系'
        '（类似 NLP 句子对用 [SEP] 拼接做匹配）。')
    doc.add_paragraph('**要点与局限**：(1) 100 个 N 仅为分隔标记，无生物学对应物——co-embedding 是**计算手段**，'
        '不代表序列真实邻接；(2) Genos 为自回归因果模型，拼接时启动子能看到增强子上下文，反之则否（单向），'
        '可尝试双向顺序；3) 我们诊断发现 embedding 两两余弦约 0.99，说明被全局组成主导，故需降维。')

    doc.add_heading('3.3 降维（KernelPCA-100）',level=2)
    doc.add_paragraph('由于 embedding 高度共线（余弦≈0.99，被 GC/长度/组成等全局信号主导），互作特异信号被淹没。'
        '我们用 **KernelPCA(RBF, 100 维)** 对 embedding 降维去噪，保留互补的序列信息。'
        '降维器**仅在训练折内拟合**（无泄漏），应用于训练与验证折。'
        '要点：降维维度经网格搜索确定（100 维，位于取样内部非边界）；也对比了 PCA、TruncatedSVD、PLS、LDA 等，'
        '其中无监督 KernelPCA 最优；有监督降维（PLS/LDA）因过拟合训练标签而不如无监督。')

    doc.add_heading('3.4 特征融合（拼接方式）',level=2)
    doc.add_paragraph('共比较 9 种配置：(1) ABC；(2) ABC+co-embedding(原始)；(3) ABC+co-embedding(RED)；'
        '(4) ABC+embedding(原始)；(5) ABC+embedding(RED)；(6) co-embedding；(7) embedding；(8) co-embedding(RED)；'
        '(9) embedding(RED)。融合方式均为**按列拼接**特征矩阵后送入分类器。'
        '要点：RED 指先降维再拼接；原始高维 embedding 直接拼接会稀释 ABC 信号（噪声过大），降维后可避免。')

    doc.add_heading('3.5 分类器',level=2)
    doc.add_paragraph('梯度提升（GB，n_estimators=300, max_depth=4, lr=0.1）与随机森林（RF，n_estimators=300）。'
        '选择理由：树模型能利用降维后的 embedding 特征，且比线性模型更适合高维稀疏信号。')

    doc.add_heading('3.6 评估与显著性',level=2)
    doc.add_paragraph('评估：**按增强子身份分组 5 折交叉验证**——同一增强子（及其全部正负配对）不跨折，避免因增强子'
        '复用导致的泄漏；降维与标准化仅在训练折拟合。报告 AUC、PR-AUC、准确率、F1。'
        '显著性：对关键对比做**配对 bootstrap**（10,000 次重采样）计算 AUC 差异的 95% 置信区间与单侧 p 值，'
        '并用配对检验验证准确率差异。')

    doc.add_heading('4. 结果',level=1)
    doc.add_heading('4.1 九种配置对比',level=2)
    tbl=doc.add_table(rows=1,cols=7); tbl.style='Light Grid Accent 1'
    for i,h in enumerate(['配置','GB-AUC','GB-PRAUC','GB-Acc','RF-AUC','RF-PRAUC','RF-Acc']): tbl.rows[0].cells[i].text=h
    cfgs=sorted(set(v['config'] for v in d.values()), key=lambda c:-max(v['auc_m'] for v in d.values() if v['config']==c))
    for c in cfgs:
        gb=next(v for v in d.values() if v['config']==c and v['model']=='GB'); rf=next(v for v in d.values() if v['config']==c and v['model']=='RF')
        r=tbl.add_row().cells
        r[0].text=c; r[1].text=f"{gb['auc_m']:.4f}"; r[2].text=f"{gb['ap_m']:.4f}"; r[3].text=f"{gb['acc_m']:.4f}"
        r[4].text=f"{rf['auc_m']:.4f}"; r[5].text=f"{rf['ap_m']:.4f}"; r[6].text=f"{rf['acc_m']:.4f}"
    doc.add_picture(os.path.join(RES,'fig_config_compare.png'),width=Inches(6.0)); doc.add_paragraph('图 1. 各配置 AUC（GB，分组 CV）。')

    doc.add_heading('4.2 ROC 与 PR',level=2)
    doc.add_picture(os.path.join(RES,'fig_roc.png'),width=Inches(3.8)); doc.add_picture(os.path.join(RES,'fig_pr.png'),width=Inches(3.8))
    doc.add_paragraph('图 2-3. 关键配置的 ROC 与 PR 曲线。')
    doc.add_heading('4.3 混淆矩阵',level=2)
    doc.add_picture(os.path.join(RES,'fig_confusion.png'),width=Inches(3.2))

    doc.add_heading('4.4 显著性检验',level=2)
    t2=doc.add_table(rows=1,cols=5); t2.style='Light Grid Accent 1'
    for i,h in enumerate(['对比','ΔAUC','95% CI','p(AUC)','p(Acc)']): t2.rows[0].cells[i].text=h
    for k,lab in [('ABC+co__vs__ABC','ABC+coembed(raw) vs ABC'),('ABC+coRED__vs__ABC','ABC+coembed(RED) vs ABC'),('ABC+co__vs__ABC+coRED','coembed(raw) vs coembed(RED)')]:
        v=sig[k]; r=t2.add_row().cells
        r[0].text=lab; r[1].text=f"{v['aucA']-v['aucB']:+.4f}"; r[2].text=f"[{v['auc_ci'][0]:+.4f},{v['auc_ci'][1]:+.4f}]"
        r[3].text=f"{v['auc_p']:.4f}"; r[4].text=f"{v['acc_p']:.4f}"
    doc.add_picture(os.path.join(RES,'fig_significance_forest.png'),width=Inches(5.5))
    doc.add_paragraph('图 4. AUC 差异森林图。')
    doc.add_paragraph('结论：ABC+co-embedding（原始或降维）均显著优于纯 ABC（p<0.001）；原始与降维差异不显著（p=0.11）。')

    doc.add_heading('4.5 特征重要性',level=2)
    doc.add_picture(os.path.join(RES,'fig_feature_importance.png'),width=Inches(4.0))
    doc.add_picture(os.path.join(RES,'fig_importance_share.png'),width=Inches(3.0))
    doc.add_paragraph('图 5-6. 最佳配置特征重要性（ABC 47.2% / coembedding 52.8%）。')

    doc.add_heading('5. 讨论（为什么这样做、关键要点）',level=1)
    for txt in [
        '序列携带互作信息：ABC（0.69）与 ABC+coembed（0.725）均远超随机，证明 sequence-only 能预测顺式 E–P 互作。',
        '距离/接触/活性是主导：ABC 的 Hi-C 接触与距离项贡献最大，说明顺式互作主要由基因组邻近性决定。',
        '深度 embedding 增量有限但真实：单独 embedding 无信号（0.53–0.56），与 ABC 融合后显著提升（+0.03）——'
        '说明 embedding 捕捉的是 ABC 之外的互补序列信息，但本身不足以单独预测互作。',
        'co-embedding 优于独立编码：让模型同时看到两个元件（同上下文）比各自独立编码更强，是更优的序列利用方式。',
        '降维的必要性：embedding 高度共线，直接拼接高维会稀释信号；KernelPCA-100 去噪后保留互补信息。',
        '分组 CV 的重要性：若不按增强子分组切分，同增强子跨折会导致泄漏、评估虚高（我们用分组 CV 纠正了单次切分的假阳性）。',
    ]: doc.add_paragraph('• '+txt)

    doc.add_heading('6. 结论',level=1)
    doc.add_paragraph('sequence-only 模型能显著预测 K562 顺式 E–P 互作（AUC 0.725，显著优于随机与纯 ABC）。'
        '最佳方式为 ABC 特征 + co-embedding 融合（GB 分类器）。序列内容携带真实但有限的信息，'
        '更大提升需引入染色质状态与三维基因组信息。')

    doc.add_heading('7. 局限与展望',level=1)
    doc.add_paragraph('(1) "负样本"为无实验证据，非证实不互作；(2) 仅顺式，trans 臂待分析；'
        '(3) embedding 增量有限，突破需染色质/3D 信息；(4) Genos 为自回归单向模型，双向编码或逐 token 注意力'
        '或可进一步利用；(5) 全量 18 组逐样本数据收集与更多分层分析仍在进行，将补充终极报告。')

    rep=os.path.join(OUT,'K562_cis_完整报告_方法详述.docx'); doc.save(rep)
    print('完整报告(方法详述)已生成:',rep)

if __name__=='__main__': main()
